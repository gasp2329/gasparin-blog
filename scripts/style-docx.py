"""
Post-process the pandoc-generated .docx to apply table styling,
status badge colors, and other formatting pandoc can't handle via reference doc.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT
import sys
import re

AMBER = RGBColor(0xB4, 0x53, 0x09)
STONE_900 = RGBColor(0x1C, 0x19, 0x17)
STONE_500 = RGBColor(0x78, 0x71, 0x6C)
STONE_200 = RGBColor(0xE7, 0xE5, 0xE4)
GREEN_BG = "DCFCE7"
GREEN_FG = RGBColor(0x16, 0x65, 0x34)
YELLOW_BG = "FEF9C3"
YELLOW_FG = RGBColor(0x85, 0x4D, 0x0E)
RED_BG = "FEE2E2"
RED_FG = RGBColor(0x99, 0x1B, 0x1B)
HEADER_BG = "F1F0EF"

docx_path = sys.argv[1] if len(sys.argv) > 1 else "exports/squad-operating-model.docx"
doc = Document(docx_path)

# ── Style all tables ──
for table in doc.tables:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Style header row
    if table.rows:
        for cell in table.rows[0].cells:
            tc = cell._element
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
                tc.insert(0, tcPr)
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{HEADER_BG}"/>'
            )
            # Remove existing shading
            for old in tcPr.findall(qn('w:shd')):
                tcPr.remove(old)
            tcPr.append(shading)

            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7)
                    run.font.bold = True
                    run.font.name = 'Calibri'
                    run.font.color.rgb = STONE_900

    # Style data rows
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if row_idx > 0:  # data row
                        run.font.size = Pt(7.5)
                        run.font.name = 'Calibri'

            # Add cell borders
            tc = cell._element
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
                tc.insert(0, tcPr)
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="2" w:space="0" w:color="E7E5E4"/>'
                f'  <w:bottom w:val="single" w:sz="2" w:space="0" w:color="E7E5E4"/>'
                f'  <w:left w:val="single" w:sz="2" w:space="0" w:color="E7E5E4"/>'
                f'  <w:right w:val="single" w:sz="2" w:space="0" w:color="E7E5E4"/>'
                f'</w:tcBorders>'
            )
            for old in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(old)
            tcPr.append(borders)

# ── Color status badge text ──
for p in doc.paragraphs:
    for run in p.runs:
        text = run.text.strip()
        if text == 'On Track':
            run.font.color.rgb = GREEN_FG
            run.font.bold = True
            run.font.size = Pt(7)
        elif text == 'At Risk':
            run.font.color.rgb = YELLOW_FG
            run.font.bold = True
            run.font.size = Pt(7)
        elif text == 'Off Track':
            run.font.color.rgb = RED_FG
            run.font.bold = True
            run.font.size = Pt(7)

# Also check inside tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    text = run.text.strip()
                    if text == 'On Track':
                        run.font.color.rgb = GREEN_FG
                        run.font.bold = True
                        run.font.size = Pt(7)
                    elif text == 'At Risk':
                        run.font.color.rgb = YELLOW_FG
                        run.font.bold = True
                        run.font.size = Pt(7)
                    elif text == 'Off Track':
                        run.font.color.rgb = RED_FG
                        run.font.bold = True
                        run.font.size = Pt(7)

# ── Make blockquotes italic with amber left border ──
for p in doc.paragraphs:
    if p.style and p.style.name in ('Block Text', 'Quote', 'IntenseQuote', 'block-quote'):
        for run in p.runs:
            run.font.italic = True
            run.font.color.rgb = STONE_500
            run.font.size = Pt(10.5)

doc.save(docx_path)
print(f'✓ Post-processed: {docx_path}')
