"""Verify styling in the generated Word document."""
from docx import Document
from docx.oxml.ns import qn
import sys

path = sys.argv[1] if len(sys.argv) > 1 else "exports/squad-operating-model.docx"
doc = Document(path)

print("=== HEADING SAMPLES ===")
count = 0
for p in doc.paragraphs:
    if p.style and p.style.name.startswith("Heading"):
        if count < 6:
            font = p.style.font
            color = font.color.rgb if font.color and font.color.rgb else "inherit"
            print(f'  [{p.style.name}] "{p.text[:60]}" | font={font.name} size={font.size} color={color} bold={font.bold}')
            count += 1

print("\n=== BODY TEXT STYLE ===")
for p in doc.paragraphs[:30]:
    if p.style and p.style.name == "Normal" and p.text.strip():
        font = p.style.font
        color = font.color.rgb if font.color and font.color.rgb else "inherit"
        print(f"  font={font.name} size={font.size} color={color}")
        break

print(f"\n=== TABLES: {len(doc.tables)} found ===")
if doc.tables:
    t = doc.tables[0]
    print(f"  First table: {len(t.rows)} rows x {len(t.columns)} cols")
    if t.rows:
        for c in t.rows[0].cells:
            for r in c.paragraphs[0].runs[:1]:
                print(f'    Header cell: "{r.text[:30]}" bold={r.font.bold} size={r.font.size}')

for section in doc.sections:
    orient = section.orientation
    w, h = section.page_width, section.page_height
    print(f"\n=== PAGE: orient={orient} width={w} height={h} ===")
    print(f"  Margins: top={section.top_margin} bottom={section.bottom_margin} left={section.left_margin} right={section.right_margin}")
